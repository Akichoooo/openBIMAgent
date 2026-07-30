"""将 开题报告.md 转换为带架构图的 Word 文档."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

MD_PATH = Path("开题报告.md")
OUTPUT_PATH = Path("开题报告.docx")
IMAGE_PATH = Path("architecture.png")


def set_run_font(run, font_name: str, size_pt: int, bold: bool = False, color: RGBColor | None = None):
    font = run.font
    font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    font.size = Pt(size_pt)
    font.bold = bold
    if color:
        font.color.rgb = color


def add_heading(doc: Document, text: str, level: int):
    """添加标题，设置中文字体."""
    para = doc.add_heading(level=level)
    run = para.add_run(text)
    set_run_font(run, "黑体", 18 - (level - 1) * 2, bold=True)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return para


def add_paragraph(doc: Document, text: str, first_line_indent: bool = True):
    """添加正文段落，支持 **加粗** 和内联代码."""
    para = doc.add_paragraph()
    if first_line_indent:
        para.paragraph_format.first_line_indent = Inches(0.3)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    para.paragraph_format.space_after = Pt(6)

    # 解析 **加粗** 和 `代码`
    parts = re.split(r"(\*\*.*?\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            set_run_font(run, "宋体", 12, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            set_run_font(run, "Courier New", 11)
        else:
            run = para.add_run(part)
            set_run_font(run, "宋体", 12)
    return para


def add_bullet(doc: Document, text: str, level: int = 0):
    """添加列表项."""
    para = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    # 解析加粗
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            set_run_font(run, "宋体", 12, bold=True)
        else:
            run = para.add_run(part)
            set_run_font(run, "宋体", 12)
    return para


def add_numbered(doc: Document, text: str):
    """添加编号列表项."""
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = para.add_run(text)
    set_run_font(run, "宋体", 12)
    return para


def add_code_block(doc: Document, lines: list[str]):
    """添加代码块."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.shading = None
    for i, line in enumerate(lines):
        if i > 0:
            para.add_run().add_break()
        run = para.add_run(line)
        set_run_font(run, "Courier New", 10)
    return para


def parse_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    """解析 markdown 表格."""
    headers = [cell.strip() for cell in lines[0].split("|")[1:-1]]
    rows = []
    for line in lines[2:]:
        cells = [cell.strip() for cell in line.split("|")[1:-1]]
        rows.append(cells)
    return headers, rows


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    """添加表格."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                set_run_font(r, "黑体", 11, bold=True)
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = cell
            for p in row_cells[i].paragraphs:
                for r in p.runs:
                    set_run_font(r, "宋体", 11)
    doc.add_paragraph()  # 表格后空一行
    return table


def add_image(doc: Document, image_path: Path, caption: str):
    """添加图片和题注."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(6.5))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run_font(run, "宋体", 10)
    return para


def main():
    doc = Document()
    # 页面设置 A4
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code_block = False
    code_lines: list[str] = []
    table_lines: list[str] = []
    in_table = False
    skip_mermaid = False

    while i < len(lines):
        line = lines[i]

        # 代码块处理
        if line.startswith("```"):
            if in_code_block:
                # 代码块结束
                if not skip_mermaid:
                    add_code_block(doc, code_lines)
                in_code_block = False
                code_lines = []
                skip_mermaid = False
            else:
                # 代码块开始
                in_code_block = True
                lang = line[3:].strip().lower()
                if lang == "mermaid":
                    skip_mermaid = True
                    # 下一张图就是架构图
            i += 1
            continue

        if in_code_block:
            if not skip_mermaid:
                code_lines.append(line)
            i += 1
            continue

        # 表格处理
        if line.startswith("|") and line.endswith("|"):
            table_lines.append(line)
            in_table = True
            i += 1
            continue
        elif in_table:
            headers, rows = parse_table(table_lines)
            add_table(doc, headers, rows)
            table_lines = []
            in_table = False
            continue

        # 标题
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line[level:].strip()
            add_heading(doc, text, level)
            if "系统总体架构图" in text:
                add_image(doc, IMAGE_PATH, "图 1  openBIMAgent 系统总体架构")
            i += 1
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 列表项
        if line.lstrip().startswith("- "):
            text = line.lstrip()[2:]
            level = (len(line) - len(line.lstrip())) // 2
            add_bullet(doc, text, level)
            i += 1
            continue

        if re.match(r"^\d+\.\s", line.lstrip()):
            text = re.sub(r"^\d+\.\s", "", line.lstrip())
            add_numbered(doc, text)
            i += 1
            continue

        # 普通段落
        add_paragraph(doc, line)
        i += 1

    # 处理末尾表格
    if table_lines:
        headers, rows = parse_table(table_lines)
        add_table(doc, headers, rows)

    doc.save(OUTPUT_PATH)
    print(f"Saved Word document: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
